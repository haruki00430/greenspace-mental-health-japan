# -*- coding: utf-8 -*-
"""Zenodo DOI を v1.0.3 (10.5281/zenodo.20986298) に一括更新（DOCX）。"""

from pathlib import Path

from docx import Document

PKG = Path(__file__).parent / "submission_package_IJEHR"
OLD_DOIS = (
    "10.5281/zenodo.20951145",
    "10.5281/zenodo.20950806",
)
NEW_DOI = "10.5281/zenodo.20986298"

DOCX_FILES = [
    PKG / "Manuscript_IJEHR_20260628.docx",
    PKG / "Manuscript_IJEHR_anonymous_20260628.docx",
    PKG / "CoverLetter_IJEHR_20260628.docx",
]


def replace_doi(text: str) -> str:
    for old in OLD_DOIS:
        text = text.replace(old, NEW_DOI)
    return text


def fix_docx(path: Path) -> int:
    doc = Document(path)
    changes = 0
    for para in doc.paragraphs:
        new_text = replace_doi(para.text)
        if new_text != para.text:
            para.text = new_text
            changes += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                new_text = replace_doi(cell.text)
                if new_text != cell.text:
                    cell.text = new_text
                    changes += 1
    if changes:
        doc.save(path)
    return changes


def main() -> None:
    total = 0
    for path in DOCX_FILES:
        if not path.exists():
            print(f"SKIP (not found): {path.name}")
            continue
        n = fix_docx(path)
        print(f"OK {path.name}: {n} block(s) updated")
        total += n
    print(f"Done. Total blocks updated: {total}")


if __name__ == "__main__":
    main()
