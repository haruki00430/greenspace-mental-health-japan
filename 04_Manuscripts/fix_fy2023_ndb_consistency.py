# -*- coding: utf-8 -*-
"""NDB 第10回＝FY 2023 への年度表記一括修正（DOCX）。"""

import re
import zipfile
from pathlib import Path

from docx import Document

PKG = Path(__file__).parent / "submission_package_IJEHR"
DOCX_FILES = [
    PKG / "Manuscript_IJEHR_20260628.docx",
    PKG / "Manuscript_IJEHR_anonymous_20260628.docx",
    PKG / "CoverLetter_IJEHR_20260628.docx",
]

REPLACEMENTS = [
    (
        "Open Data No.10 (fiscal year 2020)",
        "Open Data No.10 (fiscal year 2023)",
    ),
    (
        "National Database (NDB) Open Data (FY 2020)",
        "National Database (NDB) Open Data (FY 2023)",
    ),
    (
        "National Database of Health Insurance Claims (NDB) Open Data (FY 2020)",
        "National Database of Health Insurance Claims (NDB) Open Data (FY 2023)",
    ),
    (
        "NDB Open Data No.10 (fiscal year 2020)",
        "NDB Open Data No.10 (fiscal year 2023)",
    ),
    (
        "The study period was fiscal year 2020.",
        (
            "Prescription outcomes were measured for fiscal year 2023; socioeconomic "
            "covariates were obtained from publicly available 2020 prefecture-level "
            "statistics (see covariate list below)."
        ),
    ),
    (
        "and expressed them as tablets per 100,000 population. This outcome was interpreted",
        (
            "and expressed them as tablets per 100,000 population using prefecture "
            "population from the 2020 Population Census as the denominator. This outcome was interpreted"
        ),
    ),
    (
        "Finally, the cross-sectional design and temporal mismatch between environmental and claims data prevent causal inference.",
        (
            "Finally, the cross-sectional design, temporal mismatch between land-use and "
            "claims data, and alignment of fiscal year 2023 prescription outcomes with "
            "2020 socioeconomic covariates prevent causal inference."
        ),
    ),
]


def replace_in_text(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def fix_docx(path: Path) -> int:
    doc = Document(path)
    changes = 0
    for para in doc.paragraphs:
        new_text = replace_in_text(para.text)
        if new_text != para.text:
            para.text = new_text
            changes += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                new_text = replace_in_text(cell.text)
                if new_text != cell.text:
                    cell.text = new_text
                    changes += 1
    doc.save(path)
    sanitize_docx_package(path)
    return changes


def sanitize_docx_package(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        files = {name: zin.read(name) for name in names}
    core = files["docProps/core.xml"].decode("utf-8")
    for tag in ("dc:creator", "cp:lastModifiedBy", "dc:title", "dc:subject"):
        core = re.sub(rf"<{tag}[^>]*>.*?</{tag}>", f"<{tag}/>", core, flags=re.DOTALL)
    files["docProps/core.xml"] = core.encode("utf-8")
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, files[name])
    tmp.replace(path)


def main() -> None:
    for path in DOCX_FILES:
        if not path.exists():
            print(f"SKIP (missing): {path.name}")
            continue
        n = fix_docx(path)
        print(f"Updated {path.name}: {n} paragraph/cell blocks changed")


if __name__ == "__main__":
    main()
