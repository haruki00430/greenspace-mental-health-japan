# -*- coding: utf-8 -*-
"""IJEHR 投稿用 DOCX を submission_package_IJEHR に配置する。"""

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).parent
SRC = ROOT / "Manuscript_IJEHR.docx"
SUPP_SRC = ROOT / "Supplementary_Tables_IJEHR.docx"
PKG = ROOT / "submission_package_IJEHR"
DATE_TAG = "20260628"
MANUSCRIPT_OUT = PKG / f"Manuscript_IJEHR_{DATE_TAG}.docx"
COVER_OUT = PKG / f"CoverLetter_IJEHR_{DATE_TAG}.docx"
TITLE_OUT = PKG / f"TitlePage_IJEHR_{DATE_TAG}.docx"
DECL_OUT = PKG / f"Declarations_IJEHR_{DATE_TAG}.docx"
BIOS_OUT = PKG / f"Author_Bios_IJEHR_{DATE_TAG}.docx"
SUPP_OUT = PKG / f"Supplementary_Tables_IJEHR_{DATE_TAG}.docx"


def sanitize_docx_package(path: Path) -> None:
    """Word コアメタデータの削除と comments.xml の除去。"""
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        files = {name: zin.read(name) for name in names}

    core = files["docProps/core.xml"].decode("utf-8")
    for tag in ("dc:creator", "cp:lastModifiedBy", "dc:title", "dc:subject"):
        core = re.sub(rf"<{tag}[^>]*>.*?</{tag}>", f"<{tag}/>", core, flags=re.DOTALL)
    files["docProps/core.xml"] = core.encode("utf-8")

    drop = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/_rels/comments.xml.rels",
    }
    names = [n for n in names if n not in drop]
    files = {k: v for k, v in files.items() if k not in drop}

    rels_key = "word/_rels/document.xml.rels"
    if rels_key in files:
        rels = files[rels_key].decode("utf-8")
        rels = re.sub(
            r'<Relationship[^>]*Type="[^"]*/comments"[^>]*/>\s*',
            "",
            rels,
        )
        files[rels_key] = rels.encode("utf-8")

    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, files[name])
    tmp.replace(path)


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _add_bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _add_numbered(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _save_aux_docx(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    sanitize_docx_package(path)


def build_cover_letter_docx(path: Path) -> None:
    """IJEHR 向けカバーレター DOCX を生成する。"""
    doc = Document()
    _set_default_style(doc)

    _add_para(doc, "Date: June 28, 2026", bold=True)
    doc.add_paragraph()
    _add_para(
        doc,
        "To: The Editors, International Journal of Environmental Health Research",
        bold=True,
    )
    _add_para(
        doc,
        'Re: Submission of Research Article — "Not All Environmental '
        'Indicators Travel: Lessons from Greenspace and Mental Health Service Use in Japan"',
        bold=True,
    )
    doc.add_paragraph()
    _add_para(doc, "Dear Editor,")
    doc.add_paragraph()
    _add_para(
        doc,
        'We are pleased to submit our manuscript entitled "Not All Environmental '
        "Indicators Travel: Lessons from Greenspace and Mental Health Service Use in Japan\" "
        "for consideration as a Research Article in International Journal of Environmental "
        "Health Research.",
    )
    doc.add_paragraph()
    _add_para(doc, "Summary", bold=True)
    _add_para(
        doc,
        "This nationwide spatial ecological study examines whether prefecture-level greenspace "
        "coverage — a widely used administrative environmental indicator — predicts psychiatric "
        "medication prescription rates across all 47 Japanese prefectures. Using Moran's I and "
        "spatial regression models (OLS, Spatial Lag, Spatial Error, and Spatial Durbin) with "
        "National Database of Health Insurance Claims (NDB) Open Data (FY 2023), national "
        "land-use statistics, and socioeconomic indicators, we found that greenspace ratio did "
        "not significantly predict prescription volume after adjustment, despite significant "
        "spatial clustering in both variables. Stratified and interaction analyses showed no "
        "statistically meaningful urban-rural effect heterogeneity.",
    )
    doc.add_paragraph()
    _add_para(
        doc,
        "We interpret these findings as a proxy-validation result: administrative greenspace "
        "quantity may be too coarse a proxy for exposure pathways relevant to mental health "
        "service use, particularly when outcomes are claims-based healthcare measures rather "
        "than direct prevalence indicators.",
    )
    doc.add_paragraph()
    _add_para(doc, "Relevance to IJEHR", bold=True)
    _add_para(
        doc,
        "We believe this manuscript aligns with the journal's focus on environmental health "
        "research because:",
    )
    _add_numbered(
        doc,
        "It evaluates an administrative environmental exposure indicator (greenspace coverage) "
        "at the population level using nationally comprehensive Japanese data.",
    )
    _add_numbered(
        doc,
        "It applies spatial epidemiological methods to assess environmental-health associations "
        "while accounting for spatial autocorrelation.",
    )
    _add_numbered(
        doc,
        "The boundary-condition findings offer methodological guidance for environmental health "
        "surveillance using open administrative data.",
    )
    doc.add_paragraph()
    _add_para(doc, "Declarations", bold=True)
    _add_bullet(
        doc,
        "Ethics: Publicly available, de-identified, aggregate-level data only. Ethics committee "
        "approval was not required.",
    )
    _add_bullet(doc, "Competing interests: The authors report there are no competing interests.")
    _add_bullet(
        doc,
        "Funding: This research received no specific grant from any funding agency.",
    )
    _add_bullet(
        doc,
        "Data availability: NDB Open Data, MLIT National Land Numerical Information, e-Stat. "
        "Code and data: https://github.com/haruki00430/greenspace-mental-health-japan and "
        "https://doi.org/10.5281/zenodo.20951145.",
    )
    _add_bullet(
        doc,
        "AI use: AI-assisted tools were used for manuscript drafting and analysis scripting; "
        "all analyses and interpretations were verified by the authors.",
    )
    doc.add_paragraph()
    _add_para(
        doc,
        "This manuscript was previously submitted to Community Mental Health Journal and was "
        "declined without peer review on 28 June 2026. It has not been published elsewhere and "
        "is not under consideration by any other journal.",
    )
    doc.add_paragraph()
    _add_para(
        doc,
        "We would be grateful if you would consider our manuscript for peer review.",
    )
    doc.add_paragraph()
    _add_para(doc, "Sincerely,")
    doc.add_paragraph()
    _add_para(doc, "Haruki Saito")
    _add_para(
        doc,
        "Department of Epidemiology, Fukushima Medical University School of Medicine",
    )
    _add_para(doc, "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan")
    _add_para(doc, "m211039@fmu.ac.jp")
    _add_para(doc, "ORCID: 0009-0009-7890-6068")
    doc.add_paragraph()
    _add_para(doc, "On behalf of all authors.")

    _save_aux_docx(doc, path)


def build_title_page_docx(path: Path) -> None:
    """IJEHR タイトルページ DOCX を生成する。"""
    doc = Document()
    _set_default_style(doc)

    _add_heading(doc, "Title Page — International Journal of Environmental Health Research", level=1)
    _add_para(
        doc,
        "Reference document for Taylor & Francis Submission Portal (included as first page of "
        "manuscript).",
    )
    doc.add_paragraph()
    _add_heading(doc, "Full Title", level=2)
    _add_para(
        doc,
        "Not All Environmental Indicators Travel: Lessons from Greenspace and Mental Health "
        "Service Use in Japan",
    )
    _add_heading(doc, "Short Title", level=2)
    _add_para(doc, "Greenspace as a Proxy for Mental Health Service Use in Japan")
    doc.add_paragraph()
    _add_heading(doc, "Author 1 (Corresponding Author)", level=2)
    for line in [
        "Name: Haruki Saito",
        "Affiliation: Department of Epidemiology, Fukushima Medical University School of Medicine",
        "Address: 1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan",
        "Email: m211039@fmu.ac.jp",
        "ORCID: 0009-0009-7890-6068",
        "Role: Corresponding Author",
    ]:
        _add_para(doc, line)
    doc.add_paragraph()
    _add_heading(doc, "Author 2", level=2)
    for line in [
        "Name: Tetsuya Ohira",
        "Affiliation 1: Department of Epidemiology, Fukushima Medical University School of "
        "Medicine, Fukushima, Japan",
        "Affiliation 2: Radiation Medical Science Center for the Fukushima Health Management "
        "Survey, Fukushima Medical University, Fukushima, Japan",
        "Address: 1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan",
        "Email: teoohira@fmu.ac.jp",
        "ORCID: 0000-0003-4532-7165",
    ]:
        _add_para(doc, line)
    doc.add_paragraph()
    _add_heading(doc, "Authors' Contributions (CRediT)", level=2)
    _add_para(
        doc,
        "Haruki Saito: Conceptualization, Data curation, Formal analysis, Investigation, "
        "Methodology, Software, Visualization, Writing – original draft, Writing – review and "
        "editing.",
    )
    _add_para(
        doc,
        "Tetsuya Ohira: Conceptualization, Supervision, Writing – review and editing.",
    )
    doc.add_paragraph()
    _add_heading(doc, "Submission Information", level=2)
    for line in [
        "Journal: International Journal of Environmental Health Research (Taylor & Francis)",
        "Article type: Research Article",
        "Submission date: 2026-06-28",
        "Word count (main text only): approximately 2,490",
        "Abstract word count: approximately 159",
        "Tables and figures (main text): 6 (3 tables + 3 figures)",
        "Supplementary tables: 2 (Table S1–S2)",
        "References: 28",
        "Keywords: Greenspace; Mental health; Spatial epidemiology; Environmental epidemiology; Japan",
    ]:
        _add_para(doc, line)

    _save_aux_docx(doc, path)


def build_author_bios_docx(path: Path) -> None:
    """各著者の Biographical note DOCX を生成する（IJEHR 必須、各 ≤200 語）。"""
    doc = Document()
    _set_default_style(doc)

    _add_heading(doc, "Biographical Notes — International Journal of Environmental Health Research", level=1)
    _add_para(doc, "Each note should be ≤200 words per IJEHR author guidelines.")
    doc.add_paragraph()

    _add_heading(doc, "Haruki Saito", level=2)
    _add_para(
        doc,
        "Haruki Saito is a researcher in the Department of Epidemiology, Fukushima Medical "
        "University School of Medicine, Japan. His research focuses on spatial epidemiology, "
        "environmental health indicators, and the use of open administrative health data for "
        "population-level surveillance. He has developed analytical pipelines for National "
        "Database (NDB) Open Data and applies spatial regression methods to examine regional "
        "health disparities and proxy validity of environmental exposure measures.",
    )
    doc.add_paragraph()

    _add_heading(doc, "Tetsuya Ohira", level=2)
    _add_para(
        doc,
        "Tetsuya Ohira is Professor in the Department of Epidemiology, Fukushima Medical "
        "University School of Medicine, and affiliated with the Radiation Medical Science Center "
        "for the Fukushima Health Management Survey. His research spans environmental "
        "epidemiology, disaster and radiation health, and chronic disease prevention. He "
        "supervises population-based studies using administrative and survey data to inform "
        "public health policy in Japan.",
    )

    _save_aux_docx(doc, path)


def build_declarations_docx(path: Path) -> None:
    """IJEHR ポータル入力用 Declarations DOCX を生成する。"""
    doc = Document()
    _set_default_style(doc)

    _add_heading(doc, "Declarations — International Journal of Environmental Health Research", level=1)
    _add_para(
        doc,
        "Copy each section below into the corresponding Taylor & Francis Submission Portal fields.",
    )
    doc.add_paragraph()

    sections = [
        (
            "Ethics",
            "This study used publicly available, anonymized, aggregate-level data from government "
            "statistical databases (NDB Open Data, MLIT National Land Numerical Information L03-b "
            "land-use data, e-Stat). No individual-level data were accessed. Formal ethics "
            "committee review was not required under applicable Japanese national guidelines for "
            "secondary research using publicly available aggregate statistics. No informed consent "
            "was required.",
        ),
        ("Disclosure statement", "The authors report there are no competing interests to declare."),
        (
            "Funding",
            "This research received no specific grant from any funding agency in the public, "
            "commercial, or not-for-profit sectors.",
        ),
        (
            "Authors' Contributions (CRediT)",
            "Haruki Saito: Conceptualization, Data curation, Formal analysis, Investigation, "
            "Methodology, Software, Visualization, Writing – original draft, Writing – review and "
            "editing.\n\n"
            "Tetsuya Ohira: Conceptualization, Supervision, Writing – review and editing.\n\n"
            "All authors read and approved the final manuscript.",
        ),
        (
            "Data availability statement",
            "All data used in this study are publicly available. NDB Open Data (Ministry of Health, "
            "Labour and Welfare): https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html. "
            "National Land Numerical Information L03-b land-use data (Ministry of Land, "
            "Infrastructure, Transport and Tourism): "
            "https://nlftp.mlit.go.jp/ksj/gml/datalist/KsjTmplt-L03-b.html. e-Stat (Statistics "
            "Bureau of Japan): https://www.e-stat.go.jp/. Analysis code and aggregate "
            "prefecture-level data (N = 47) are openly available at "
            "https://github.com/haruki00430/greenspace-mental-health-japan and "
            "https://doi.org/10.5281/zenodo.20951145.",
        ),
        (
            "Declaration of generative AI",
            "During the preparation and writing of this work, the authors used AI-assisted tools "
            "to support manuscript drafting and statistical analysis scripting. Cursor 3.0 "
            "(Anysphere) and Google Antigravity (Google) were used for AI-assisted writing and "
            "Python code development. Large language models used through these platforms included "
            "Claude Sonnet 4.6 and Claude Opus 4.8 (Anthropic) and Gemini 3 Pro (Google). These "
            "tools were used only for text drafting and code generation; no generative AI or "
            "AI-assisted tools were used to create, alter, or otherwise process any figures, images, "
            "or artwork in this manuscript. The authors reviewed and edited all AI-assisted outputs "
            "and were responsible for the study design, selection of statistical methods, "
            "interpretation of findings, conclusions, and final reference list. The authors take "
            "full responsibility for the integrity and accuracy of the final content. AI was not "
            "listed as an author.",
        ),
        (
            "Acknowledgements",
            "We thank the Ministry of Health, Labour and Welfare for providing access to the "
            "National Database of Health Insurance Claims (NDB) Open Data, and the Ministry of "
            "Land, Infrastructure, Transport and Tourism for providing land-use data used to "
            "construct greenspace indicators.",
        ),
    ]

    for title, body in sections:
        _add_heading(doc, title, level=2)
        for paragraph in body.split("\n\n"):
            _add_para(doc, paragraph)
        doc.add_paragraph()

    _save_aux_docx(doc, path)


def build_manuscript_docx() -> None:
    """主原稿 DOCX を submission_package_IJEHR に配置する。"""
    PKG.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, MANUSCRIPT_OUT)
    sanitize_docx_package(MANUSCRIPT_OUT)


def build_supplementary_docx() -> None:
    """Supplementary Tables DOCX を配置する。"""
    if not SUPP_SRC.exists():
        raise FileNotFoundError(
            f"Supplementary source not found: {SUPP_SRC}. "
            "Run: quarto render Supplementary_Tables_IJEHR.qmd --to docx"
        )
    shutil.copy2(SUPP_SRC, SUPP_OUT)
    sanitize_docx_package(SUPP_OUT)


def build_aux_docx_only() -> None:
    """主原稿を触らず、付属 DOCX のみ再生成する。"""
    PKG.mkdir(parents=True, exist_ok=True)
    build_cover_letter_docx(COVER_OUT)
    build_title_page_docx(TITLE_OUT)
    build_author_bios_docx(BIOS_OUT)
    build_declarations_docx(DECL_OUT)
    print(f"Wrote: {COVER_OUT.name}")
    print(f"Wrote: {TITLE_OUT.name}")
    print(f"Wrote: {BIOS_OUT.name}")
    print(f"Wrote: {DECL_OUT.name}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(
            f"Manuscript source not found: {SRC}. "
            "Run: quarto render Manuscript_IJEHR.qmd --to docx"
        )

    build_manuscript_docx()
    build_supplementary_docx()
    build_cover_letter_docx(COVER_OUT)
    build_title_page_docx(TITLE_OUT)
    build_author_bios_docx(BIOS_OUT)
    build_declarations_docx(DECL_OUT)

    print(f"Wrote: {MANUSCRIPT_OUT.name}")
    print(f"Wrote: {SUPP_OUT.name}")
    print(f"Wrote: {COVER_OUT.name}")
    print(f"Wrote: {TITLE_OUT.name}")
    print(f"Wrote: {BIOS_OUT.name}")
    print(f"Wrote: {DECL_OUT.name}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "--aux-only":
        build_aux_docx_only()
    else:
        main()
