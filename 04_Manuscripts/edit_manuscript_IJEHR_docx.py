# -*- coding: utf-8 -*-
"""Not_All_Environmental_Indicators_Travel_20260627.docx を IJEHR 投稿用に直接修正する。"""

from __future__ import annotations

import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).parent
SRC = ROOT / "Not_All_Environmental_Indicators_Travel_20260627.docx"
DATE_TAG = "20260628"
TABLE3_LEGEND = "†p<0.10, *p<0.05, **p<0.01, ***p<0.001"
SIG_CELL_RE = re.compile(r"^(.+?)(\*+)?(\s*\([^)]+\))?$")
OUT = ROOT / "submission_package_IJEHR" / f"Manuscript_IJEHR_{DATE_TAG}.docx"
SUPP_OUT = ROOT / "submission_package_IJEHR" / f"Supplementary_Tables_IJEHR_{DATE_TAG}.docx"

UNSTRUCTURED_ABSTRACT = (
    "Environmental epidemiology often relies on administrative indicators as proxies for "
    "human exposure. Greenspace quantity is widely used as a protective environmental "
    "determinant of mental health, but area-level coverage may not capture accessibility, "
    "quality, utilization, or healthcare pathways. We examined whether prefecture-level "
    "greenspace coverage predicts psychiatric medication prescriptions across all 47 Japanese "
    "prefectures using National Database (NDB) Open Data (FY 2023), national land-use "
    "statistics, and socioeconomic indicators. Spatial autocorrelation was assessed with "
    "Moran's I; ordinary least squares, spatial lag, spatial error, and Spatial Durbin "
    "models were compared. Stratified and interaction analyses tested urban-rural "
    "heterogeneity. Greenspace ratio and prescription rates showed significant spatial "
    "clustering, but after socioeconomic and spatial adjustment greenspace did not predict "
    "prescriptions (Spatial Durbin β=1,326, p=0.836). Stratified coefficients differed in "
    "direction but interaction tests showed no effect modification. Administrative greenspace "
    "quantity was a weak proxy for mental health-relevant environmental exposure; indicators "
    "should be validated against intended exposure pathways before use in ecological mental "
    "health surveillance."
)

KEYWORDS = (
    "Keywords: Greenspace; Mental health; Spatial epidemiology; "
    "Environmental epidemiology; Japan"
)

GREENSSPACE_TEXT = (
    "Greenspace data were obtained from the National Land Numerical Information land-use "
    "subdivided mesh dataset (L03-b) published by the Ministry of Land, Infrastructure, "
    "Transport and Tourism (19). Prefecture-level forest and park areas were aggregated from "
    "1-km mesh cells and divided by total prefecture area (km²), expressed as a percentage. "
    "Forest area included the forest land-use class; park area included parks and green space. "
    "This variable was treated as an administrative environmental proxy: it captures quantity, "
    "but not accessibility, quality, safety, biodiversity, perceived restorative value, or "
    "actual use."
)

COVARIATE_HEADING = "Socioeconomic and Healthcare Covariates"

COVARIATE_INTRO = (
    "Covariates were obtained from the e-Stat portal (Statistics Bureau of Japan) (20) and "
    "the Ministry of Health, Labour and Welfare medical facility survey:"
)

COVARIATE_BULLETS = [
    "Aging rate (%): Proportion of population aged ≥65 years (Population Census 2020)",
    "Unemployment rate (%): Complete unemployment rate (Population Census 2020)",
    "Population density (persons/km²): Total population divided by habitable land area",
    "Income per capita (JPY): Per-capita taxable income (Municipal Taxation Statistics, FY 2020)",
    "College graduate rate (%): Proportion of population aged ≥25 years with a college degree or higher (Employment Status Survey 2017)",
    "Psychiatric clinic density: Number of psychiatric clinics per 100,000 population (Medical Facility Survey 2020)",
]

COVARIATE_CLOSING = (
    "These variables were selected a priori as potential confounders based on previous "
    "literature documenting associations between socioeconomic status, healthcare access, "
    "and both greenspace availability and mental health outcomes (21,22)."
)

ETHICS_TEXT = (
    "Ethics: This study used publicly available, anonymized, aggregate-level data from "
    "government statistical databases, including the NDB Open Data (Ministry of Health, "
    "Labour and Welfare of Japan), the National Land Numerical Information land-use dataset "
    "(Ministry of Land, Infrastructure, Transport and Tourism), and e-Stat (Statistics "
    "Bureau of Japan). No individual-level data were accessed. Under applicable Japanese "
    "national guidelines, formal ethics committee review is not required for secondary "
    "research using publicly available aggregate statistics. No informed consent was required."
)

AI_METHODS_NOTE = (
    "Use of AI tools: During analysis scripting and manuscript preparation, the authors used "
    "AI-assisted coding tools (Cursor, Claude Code). All statistical analyses were "
    "implemented and verified by the authors; NDB raw data were not transmitted to AI "
    "systems. See Declaration of generative AI below for full disclosure."
)

ACK_TEXT = (
    "We thank the Ministry of Health, Labour and Welfare for providing access to the "
    "National Database of Health Insurance Claims (NDB) Open Data, and the Ministry of "
    "Land, Infrastructure, Transport and Tourism for providing land-use data used to "
    "construct greenspace indicators."
)

DATA_AVAILABILITY = (
    "All data used in this study are publicly available. NDB Open Data (Ministry of Health, "
    "Labour and Welfare): https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html. "
    "National Land Numerical Information L03-b land-use data (Ministry of Land, Infrastructure, "
    "Transport and Tourism): https://nlftp.mlit.go.jp/ksj/gml/datalist/KsjTmplt-L03-b.html. "
    "e-Stat (Statistics Bureau of Japan): https://www.e-stat.go.jp/. Analysis code and "
    "aggregate prefecture-level data (N = 47) are openly available at "
    "https://github.com/haruki00430/greenspace-mental-health-japan and "
    "https://doi.org/10.5281/zenodo.20986298."
)

DISCLOSURE = "The authors report there are no competing interests to declare."

CREDIT = (
    "Haruki Saito: Conceptualization, Data curation, Formal analysis, Investigation, "
    "Methodology, Software, Visualization, Writing – original draft, Writing – review and "
    "editing.\n\n"
    "Tetsuya Ohira: Conceptualization, Supervision, Writing – review and editing.\n\n"
    "All authors read and approved the final manuscript."
)

MLIT_REFERENCE = (
    "19. Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical "
    "Information land-use subdivided mesh data (L03-b). "
    "https://nlftp.mlit.go.jp/ksj/gml/datalist/KsjTmplt-L03-b.html; 2022."
)

TITLE_PAGE_LINES = [
    "Haruki Saito¹,* and Tetsuya Ohira¹,²",
    "¹ Department of Epidemiology, Fukushima Medical University School of Medicine, "
    "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan",
    "² Radiation Medical Science Center for the Fukushima Health Management Survey, "
    "Fukushima Medical University, Fukushima, Japan",
    "*Corresponding author: Haruki Saito. Email: m211039@fmu.ac.jp. "
    "ORCID: 0009-0009-7890-6068",
    "Tetsuya Ohira. Email: teoohira@fmu.ac.jp. ORCID: 0000-0003-4532-7165",
    "Article type: Research Article",
    "Word count (main text): approximately 2,490 | Abstract: approximately 159 | "
    "Tables and figures: 6 (3 tables + 3 figures) | References: 28",
]


def sanitize_docx_package(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        files = {name: zin.read(name) for name in names}

    core = files["docProps/core.xml"].decode("utf-8")
    for tag in ("dc:creator", "cp:lastModifiedBy", "dc:title", "dc:subject"):
        core = re.sub(rf"<{tag}[^>]*>.*?</{tag}>", f"<{tag}/>", core, flags=re.DOTALL)
    files["docProps/core.xml"] = core.encode("utf-8")

    drop = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/_rels/comments.xml.rels",
    }
    names = [n for n in names if n not in drop]
    files = {k: v for k, v in files.items() if k not in drop}

    rels_key = "word/_rels/document.xml.rels"
    if rels_key in files:
        rels = files[rels_key].decode("utf-8")
        rels = re.sub(
            r'<Relationship[^>]*Type="[^"]*/comments"[^>]*/>\s*',
            "",
            rels,
        )
        files[rels_key] = rels.encode("utf-8")

    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, files[name])
    tmp.replace(path)


def set_para_text(para, text: str) -> None:
    """段落テキストを置換する（書式は最初の run を維持）。"""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def delete_table(table) -> None:
    element = table._element
    element.getparent().remove(element)


def find_paragraph(doc: Document, predicate) -> int | None:
    for i, para in enumerate(doc.paragraphs):
        if predicate(para):
            return i
    return None


def replace_in_paragraphs(doc: Document, old: str, new: str) -> int:
    count = 0
    for para in doc.paragraphs:
        if old in para.text:
            set_para_text(para, para.text.replace(old, new))
            count += 1
    return count


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    set_para_text(new_para, text)
    return new_para


def edit_main_manuscript(doc: Document) -> None:
    # --- Title page ---
    set_para_text(doc.paragraphs[1], TITLE_PAGE_LINES[0])
    set_para_text(doc.paragraphs[2], TITLE_PAGE_LINES[1])
    set_para_text(doc.paragraphs[3], TITLE_PAGE_LINES[2])
    anchor = doc.paragraphs[3]
    for line in TITLE_PAGE_LINES[3:]:
        anchor = insert_paragraph_after(anchor, line, "Author")

    # --- Abstract: 非構造化 1 段落 ---
    set_para_text(doc.paragraphs[6], UNSTRUCTURED_ABSTRACT)
    for idx in sorted([7, 8, 9, 10], reverse=True):
        delete_paragraph(doc.paragraphs[idx])
    set_para_text(doc.paragraphs[7], KEYWORDS)

    # --- Methods → Materials and methods（見出しのみ） ---
    for para in doc.paragraphs:
        if para.text.strip() == "Methods":
            set_para_text(para, "Materials and methods")

    # --- Greenspace data source ---
    idx = find_paragraph(doc, lambda p: p.text.startswith("Greenspace data were obtained"))
    if idx is not None:
        set_para_text(doc.paragraphs[idx], GREENSSPACE_TEXT)

    # --- Covariates ---
    idx = find_paragraph(doc, lambda p: p.text.strip() == "Socioeconomic Covariates")
    if idx is not None:
        set_para_text(doc.paragraphs[idx], COVARIATE_HEADING)

    idx = find_paragraph(
        doc,
        lambda p: p.text.startswith("Socioeconomic indicators were obtained"),
    )
    if idx is not None:
        set_para_text(doc.paragraphs[idx], COVARIATE_INTRO)
        anchor = doc.paragraphs[idx]
        # 既存 4 項目を削除
        for _ in range(4):
            nxt = anchor._element.getnext()
            if nxt is not None:
                nxt.getparent().remove(nxt)
        for bullet in COVARIATE_BULLETS:
            anchor = insert_paragraph_after(anchor, bullet, "Compact")
        insert_paragraph_after(anchor, COVARIATE_CLOSING, "First Paragraph")

    # --- Ethics & AI note before Results ---
    idx = find_paragraph(doc, lambda p: p.text.strip() == "Results")
    if idx is not None:
        ethics_para = doc.paragraphs[idx - 1]
        insert_paragraph_after(ethics_para, ETHICS_TEXT, "Body Text")
        insert_paragraph_after(doc.paragraphs[idx - 1], AI_METHODS_NOTE, "Body Text")

    # --- Results table references ---
    replace_in_paragraphs(
        doc,
        "single household rate (r=0.664, p<0.001).",
        "single household rate (r=0.664, p<0.001). "
        "The full correlation matrix is reported in Supplementary Table S1.",
    )
    replace_in_paragraphs(doc, "Table 4 summarizes", "Table 3 summarizes")
    replace_in_paragraphs(
        doc,
        "Table 5 presents stratified analysis results",
        "Supplementary Table S2 presents stratified analysis results",
    )

    # --- End matter ---
    idx = find_paragraph(doc, lambda p: p.text.strip() == "Acknowledgments")
    if idx is not None and idx + 1 < len(doc.paragraphs):
        set_para_text(doc.paragraphs[idx + 1], ACK_TEXT)

    idx = find_paragraph(doc, lambda p: p.text.strip() == "Competing Interests")
    if idx is not None:
        set_para_text(doc.paragraphs[idx], "Disclosure statement")
        if idx + 1 < len(doc.paragraphs):
            set_para_text(doc.paragraphs[idx + 1], DISCLOSURE)

    idx = find_paragraph(doc, lambda p: p.text.strip() == "Data Availability")
    if idx is not None:
        set_para_text(doc.paragraphs[idx], "Data availability statement")
        # 後続の bullet と upon request 段落を整理
        to_delete = []
        for j in range(idx + 1, min(idx + 6, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.strip()
            if t.startswith("NDB Open Data") or t.startswith("Biodiversity") or t.startswith("e-Stat") or t.startswith("Analysis code"):
                to_delete.append(j)
        for j in sorted(to_delete, reverse=True):
            delete_paragraph(doc.paragraphs[j])
        insert_paragraph_after(doc.paragraphs[idx], DATA_AVAILABILITY, "First Paragraph")

    # Authors' contributions を Funding の前に挿入
    idx = find_paragraph(doc, lambda p: p.text.strip() == "Funding")
    if idx is not None:
        h = insert_paragraph_after(doc.paragraphs[idx - 1], "Authors' contributions", "Heading 1")
        insert_paragraph_after(h, CREDIT, "First Paragraph")

    # --- References: Biodic → MLIT ---
    idx = find_paragraph(doc, lambda p: p.text.startswith("19. Biodiversity Center"))
    if idx is not None:
        set_para_text(doc.paragraphs[idx], MLIT_REFERENCE)

    # --- Tables section captions: remove Table 2 & 5, renumber ---
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("Table 2. Pearson Correlation"):
            delete_paragraph(para)
            break
    for para in doc.paragraphs:
        if para.text.strip().startswith("p < 0.05") and "0.001" in para.text:
            delete_paragraph(para)
            break
    replace_in_paragraphs(doc, "Table 3. Global Moran", "Table 2. Global Moran")
    replace_in_paragraphs(doc, "Table 4. Spatial Regression", "Table 3. Spatial Regression")

    to_remove = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("Table 5. Stratified Analysis"):
            to_remove.append(para)
        elif t == "Stratified Analysis (Separate Models by Stratum)":
            to_remove.append(para)
        elif t == "Interaction Tests (Full Sample Models, N=47)":
            to_remove.append(para)
        elif t.startswith("Note: Urban indicator is a binary"):
            to_remove.append(para)
    for para in to_remove:
        delete_paragraph(para)

    # --- Remove tables: correlation (1), stratified (4), interaction (5) ---
    if len(doc.tables) >= 6:
        delete_table(doc.tables[5])
        delete_table(doc.tables[4])
        delete_table(doc.tables[1])

    fix_table3_significance(doc)


def set_para_plain(para, text: str) -> None:
    """書式を除去してプレーンテキストに統一する。"""
    para.clear()
    run = para.add_run(text)
    run.bold = False
    run.italic = False
    run.font.superscript = False


def set_cell_plain_inline(cell, text: str) -> None:
    """係数セルを元原稿どおりインラインアスタリスクのプレーンテキストに戻す。"""
    cell.text = text.strip()


def set_cell_bold_value(cell, value: str) -> None:
    """セルをアスタリスクなしの太字値にする（AIC 最小モデルの強調用）。"""
    cell.text = value.strip()
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].bold = True


def fix_table3_aic_emphasis(tbl) -> None:
    """AIC 行 SDM 列: 1368.48 を太字（有意性アスタリスクではない）。"""
    for row in tbl.rows:
        if row.cells[0].text.strip() != "AIC":
            continue
        if len(row.cells) > 3:
            set_cell_bold_value(row.cells[3], "1368.48")
        break


def fix_table3_significance(doc: Document) -> None:
    """Table 3: QMD→DOCX 変換で化けた凡例を正し、表内はインライン * を維持する。"""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Table 3. Spatial Regression"):
            if i + 1 < len(doc.paragraphs):
                set_para_plain(doc.paragraphs[i + 1], TABLE3_LEGEND)
            break
    else:
        for para in doc.paragraphs:
            t = para.text.strip()
            if "p<0.10" in t and "p<0.001" in t and len(t) < 80:
                set_para_plain(para, TABLE3_LEGEND)
                break

    if len(doc.tables) < 3 or not SRC.exists():
        return

    base_doc = Document(SRC)
    if len(base_doc.tables) >= 4:
        src_tbl = base_doc.tables[3]
        dst_tbl = doc.tables[2]
        if len(src_tbl.rows) == len(dst_tbl.rows):
            for ri in range(len(dst_tbl.rows)):
                for ci in range(1, len(dst_tbl.rows[ri].cells)):
                    if ci < len(src_tbl.rows[ri].cells):
                        dst_tbl.rows[ri].cells[ci].text = src_tbl.rows[ri].cells[ci].text
            fix_table3_aic_emphasis(dst_tbl)


def build_supplementary_docx(main_doc: Document) -> None:
    """主原稿から除去した表 2 件で Supplementary DOCX を生成する。"""
    # 元 DOCX から correlation / stratified 表を再取得
    src_doc = Document(SRC)
    supp = Document()
    style = supp.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    supp.add_heading("Supplementary Material", level=1)
    supp.add_paragraph(
        "Not All Environmental Indicators Travel: Lessons from Greenspace and "
        "Mental Health Service Use in Japan"
    )
    supp.add_paragraph(
        "International Journal of Environmental Health Research — Supplementary Tables"
    )
    supp.add_heading(
        "Supplementary Table S1. Pearson Correlation Matrix of Study Variables (N=47)",
        level=2,
    )
    corr_tbl = src_doc.tables[1]
    new_tbl = supp.add_table(rows=len(corr_tbl.rows), cols=len(corr_tbl.columns))
    new_tbl.style = corr_tbl.style
    for r, row in enumerate(corr_tbl.rows):
        for c, cell in enumerate(row.cells):
            new_tbl.rows[r].cells[c].text = cell.text
    supp.add_paragraph("*p<0.05, **p<0.01, ***p<0.001")

    supp.add_heading(
        "Supplementary Table S2. Stratified Analysis and Interaction Tests",
        level=2,
    )
    supp.add_heading("Stratified Analysis (Separate OLS Models by Stratum)", level=3)
    strat_tbl = src_doc.tables[4]
    t2 = supp.add_table(rows=len(strat_tbl.rows), cols=len(strat_tbl.columns))
    t2.style = strat_tbl.style
    for r, row in enumerate(strat_tbl.rows):
        for c, cell in enumerate(row.cells):
            t2.rows[r].cells[c].text = cell.text

    supp.add_heading("Interaction Tests (Full Sample, N=47)", level=3)
    int_tbl = src_doc.tables[5]
    t3 = supp.add_table(rows=len(int_tbl.rows), cols=len(int_tbl.columns))
    t3.style = int_tbl.style
    for r, row in enumerate(int_tbl.rows):
        for c, cell in enumerate(row.cells):
            t3.rows[r].cells[c].text = cell.text
    supp.add_paragraph(
        "Note: Urban indicator is a binary variable (1=urban, 0=rural) based on "
        "median population density."
    )

    SUPP_OUT.parent.mkdir(parents=True, exist_ok=True)
    supp.save(SUPP_OUT)
    sanitize_docx_package(SUPP_OUT)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"Source not found: {SRC}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT)

    doc = Document(OUT)
    edit_main_manuscript(doc)
    doc.save(OUT)
    sanitize_docx_package(OUT)

    build_supplementary_docx(doc)

    print(f"Edited: {OUT}")
    print(f"Wrote: {SUPP_OUT}")


if __name__ == "__main__":
    main()
