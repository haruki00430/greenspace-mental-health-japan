# -*- coding: utf-8 -*-
"""IJEHR DOCX 修正後のクリーンアップ（残存段落・誤置換の修正）。"""

import re
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document

OUT = Path(__file__).parent / "submission_package_IJEHR" / "Manuscript_IJEHR_20260628.docx"
LEGACY_OUT = Path(__file__).parent / "submission_package_IJEHR" / "Manuscript_IJEHR_20260627.docx"

TABLE3_LEGEND = "†p<0.10, *p<0.05, **p<0.01, ***p<0.001"
SIG_CELL_RE = re.compile(r"^(.+?)(\*+)?(\s*\([^)]+\))?$")


def sanitize_docx_package(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        files = {name: zin.read(name) for name in names}
    core = files["docProps/core.xml"].decode("utf-8")
    for tag in ("dc:creator", "cp:lastModifiedBy", "dc:title", "dc:subject"):
        core = re.sub(rf"<{tag}[^>]*>.*?</{tag}>", f"<{tag}/>", core, flags=re.DOTALL)
    files["docProps/core.xml"] = core.encode("utf-8")
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, files[name])
    tmp.replace(path)


def set_para_text(para, text: str) -> None:
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def set_para_plain(para, text: str) -> None:
    """書式を除去してプレーンテキストに統一する。"""
    para.clear()
    run = para.add_run(text)
    run.bold = False
    run.italic = False
    run.font.superscript = False


def set_cell_plain_inline(cell, text: str) -> None:
    """係数セルを元原稿どおりインラインアスタリスクのプレーンテキストに戻す。"""
    cell.text = text.strip()


def set_cell_bold_value(cell, value: str) -> None:
    """セルをアスタリスクなしの太字値にする（AIC 最小モデルの強調用）。"""
    cell.text = value.strip()
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].bold = True


def fix_table3_aic_emphasis(tbl) -> None:
    """AIC 行 SDM 列: 1368.48 を太字（有意性アスタリスクではない）。"""
    for row in tbl.rows:
        if row.cells[0].text.strip() != "AIC":
            continue
        if len(row.cells) > 3:
            set_cell_bold_value(row.cells[3], "1368.48")
        break


def fix_table3_significance(doc: Document) -> None:
    """Table 3: QMD→DOCX 変換で化けた凡例を正し、表内はインライン * を維持する。"""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Table 3. Spatial Regression"):
            if i + 1 < len(doc.paragraphs):
                set_para_plain(doc.paragraphs[i + 1], TABLE3_LEGEND)
            break
    else:
        for para in doc.paragraphs:
            t = para.text.strip()
            if "p<0.10" in t and "p<0.001" in t and len(t) < 80:
                set_para_plain(para, TABLE3_LEGEND)
                break

    if len(doc.tables) < 3:
        return

    # ベース原稿（Not_All）から空間回帰表のセル文字列を復元
    base_path = Path(__file__).parent / "Not_All_Environmental_Indicators_Travel_20260627.docx"
    if base_path.exists():
        base_doc = Document(base_path)
        if len(base_doc.tables) >= 4:
            src = base_doc.tables[3]
            dst = doc.tables[2]
            if len(src.rows) == len(dst.rows) and len(src.columns) == len(dst.columns):
                for ri, row in enumerate(dst.rows):
                    for ci in range(1, len(row.cells)):
                        dst.rows[ri].cells[ci].text = src.rows[ri].cells[ci].text
                fix_table3_aic_emphasis(dst)


def fix_table3_footnote(doc: Document) -> None:
    """後方互換エイリアス。"""
    fix_table3_significance(doc)


def delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_after(paragraph, text: str, style: str = "Author"):
    from docx.text.paragraph import Paragraph

    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    set_para_text(new_para, text)
    return new_para


def main() -> None:
    import shutil

    if not OUT.exists() and LEGACY_OUT.exists():
        shutil.copy2(LEGACY_OUT, OUT)

    doc = Document(OUT)

    # 残存する構造化 Abstract 断片を削除
    drop_prefixes = (
        "Materials and methods: We conducted a nationwide spatial ecological",
        "Results: Greenspace ratio and psychiatric medication",
        "Conclusions: Administrative greenspace quantity was a weak proxy",
        "Keywords: Greenspace; Mental health; Psychiatric medication",
    )
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if any(t.startswith(p) for p in drop_prefixes):
            delete_paragraph(para)

    # Abstract 見出しを本文の前に配置
    abs_text_idx = None
    abs_head_idx = None
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t.startswith("Environmental epidemiology often relies on administrative indicators"):
            abs_text_idx = i
        if t == "Abstract":
            abs_head_idx = i
    if abs_text_idx is not None and abs_head_idx is not None and abs_head_idx > abs_text_idx:
        heading_para = doc.paragraphs[abs_head_idx]
        heading_para._element.getparent().remove(heading_para._element)
        target = doc.paragraphs[abs_text_idx]
        from docx.text.paragraph import Paragraph

        new_p = deepcopy(heading_para._element)
        target._element.addprevious(new_p)
        new_para = Paragraph(new_p, target._parent)
        new_para.style = "Heading 1"
        set_para_text(new_para, "Abstract")

    # 重複 covariate 段落を削除（旧版テキスト）
    seen_covariate_close = False
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t.startswith("These variables were selected a priori"):
            if "healthcare access" in t:
                if seen_covariate_close:
                    delete_paragraph(para)
                else:
                    seen_covariate_close = True
            elif seen_covariate_close:
                delete_paragraph(para)

    # 参考文献の誤置換を修正
    for para in doc.paragraphs:
        if "Materials and methods and Models" in para.text:
            set_para_text(
                para,
                para.text.replace(
                    "Materials and methods and Models",
                    "Methods and Models",
                ),
            )

    # Data availability の重複導入文を削除
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t == "All data used in this study are publicly available from the following sources:":
            delete_paragraph(para)

    # タイトルページに Article type 行が無ければ追加
    has_article_type = any("Article type: Research Article" in p.text for p in doc.paragraphs)
    if not has_article_type:
        for para in doc.paragraphs:
            if "ORCID: 0000-0003-4532-7165" in para.text and "Tetsuya Ohira" in para.text:
                a = insert_after(
                    para,
                    "Article type: Research Article",
                )
                insert_after(
                    a,
                    "Word count (main text): approximately 2,490 | Abstract: approximately 159 | "
                    "Tables and figures: 6 (3 tables + 3 figures) | References: 28",
                )
                break

    # メールアドレス更新
    for para in doc.paragraphs:
        if "haruki00430@gmail.com" in para.text:
            set_para_text(para, para.text.replace("haruki00430@gmail.com", "m211039@fmu.ac.jp"))

    fix_table3_significance(doc)

    doc.save(OUT)
    sanitize_docx_package(OUT)
    print(f"Cleaned: {OUT}")


if __name__ == "__main__":
    main()
