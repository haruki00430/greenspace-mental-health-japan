"""
DOCXファイルに図を挿入するスクリプト
Manuscript_greenspace_mental_health.docxに図を挿入する
"""

import os
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# プロジェクトルート
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ファイルパス
DOCX_PATH = os.path.join(PROJECT_ROOT, "04_Manuscripts", "Manuscript_greenspace_mental_health.docx")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "04_Manuscripts", "Manuscript_greenspace_mental_health_with_figures.docx")

# 図ファイルのマッピング（Figure番号 → ファイルパス）
FIGURE_FILES = {
    "Figure 1": os.path.join(PROJECT_ROOT, "results", "figures", "choropleth_greenspace.png"),
    "Figure 2": os.path.join(PROJECT_ROOT, "results", "figures", "choropleth_prescription.png"),
    "Figure 3": os.path.join(PROJECT_ROOT, "results", "figures", "lisa_map_prescription.png"),
}

def insert_paragraph_after(paragraph, text=None, style=None):
    """段落の後に新しい段落を挿入"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[-1]
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def insert_figures_to_docx(docx_path, output_path, figure_files):
    """
    DOCXファイルに図を挿入する

    Parameters
    ----------
    docx_path : str
        入力DOCXファイルパス
    output_path : str
        出力DOCXファイルパス
    figure_files : dict
        Figure番号 → ファイルパスのマッピング
    """
    print("="*80)
    print(" DOCXファイルに図を挿入")
    print("="*80)
    print()

    # DOCXファイルを読み込み
    print(f"[STEP 1] DOCXファイルを読み込み: {os.path.basename(docx_path)}")
    doc = Document(docx_path)

    # 図ファイルの存在確認
    print("\n[STEP 2] 図ファイルの存在確認")
    for fig_name, fig_path in figure_files.items():
        if os.path.exists(fig_path):
            print(f"  [OK] {fig_name}: {os.path.basename(fig_path)}")
        else:
            print(f"  [NG] {fig_name}: ファイルが見つかりません - {fig_path}")

    # 段落をスキャンしてFigure参照を見つける
    print("\n[STEP 3] Figure参照を検索して画像を挿入")
    inserted_count = 0
    paragraphs_to_process = []

    # まず処理対象の段落を特定
    for para in doc.paragraphs:
        text = para.text.strip()
        for fig_name, fig_path in figure_files.items():
            if text.startswith(fig_name) and os.path.exists(fig_path):
                paragraphs_to_process.append((para, fig_name, fig_path))
                break

    # 逆順で処理（段落追加時のインデックスずれを防ぐ）
    for para, fig_name, fig_path in reversed(paragraphs_to_process):
        # 新しい段落を追加
        new_para = insert_paragraph_after(para)
        # 画像を追加（幅6インチ）
        new_para.add_run().add_picture(fig_path, width=Inches(6.0))

        print(f"  [OK] {fig_name} を挿入しました")
        inserted_count += 1

    # ファイルを保存
    print(f"\n[STEP 4] DOCXファイルを保存: {os.path.basename(output_path)}")
    doc.save(output_path)

    print(f"\n完了: {inserted_count}個の図を挿入しました")
    print(f"出力ファイル: {output_path}")

if __name__ == "__main__":
    try:
        insert_figures_to_docx(DOCX_PATH, OUTPUT_PATH, FIGURE_FILES)
    except FileNotFoundError as e:
        print(f"エラー: ファイルが見つかりません - {e}")
    except Exception as e:
        print(f"エラー: {e}")
        import traceback
        traceback.print_exc()
