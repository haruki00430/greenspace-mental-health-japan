# -*- coding: utf-8 -*-
"""投稿用 DOCX を submission_package_CMHJ に配置する。"""

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).parent
SRC = ROOT / "Manuscript_CMHJ.docx"
PKG = ROOT / "submission_package_CMHJ"
APA_OUT = PKG / "Manuscript_CMHJ_APA_20260627.docx"
BLIND_OUT = PKG / "Manuscript_CMHJ_blinded_20260627.docx"


def clear_docx_metadata(path: Path) -> None:
    """Word コアメタデータから著者情報を削除する。"""
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        files = {name: zin.read(name) for name in names}

    core = files["docProps/core.xml"].decode("utf-8")
    for tag in ("dc:creator", "cp:lastModifiedBy", "dc:title", "dc:subject"):
        core = re.sub(rf"<{tag}[^>]*>.*?</{tag}>", f"<{tag}/>", core, flags=re.DOTALL)
    files["docProps/core.xml"] = core.encode("utf-8")

    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, files[name])
    tmp.replace(path)


def anonymize_credits(doc: Document) -> None:
    """CRediT 行をダブルブラインド用に置換する。"""
    replacements = {
        "Author 1:": "[Details removed for peer review]",
        "Author 2:": "[Details removed for peer review]",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        for old, new in replacements.items():
            if text.startswith(old):
                para.clear()
                para.add_run(new)
                break


def main() -> None:
    PKG.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, APA_OUT)
    clear_docx_metadata(APA_OUT)

    shutil.copy2(SRC, BLIND_OUT)
    doc = Document(BLIND_OUT)
    anonymize_credits(doc)
    doc.save(BLIND_OUT)
    clear_docx_metadata(BLIND_OUT)

    print(f"Wrote: {APA_OUT.name}")
    print(f"Wrote: {BLIND_OUT.name}")


if __name__ == "__main__":
    main()
